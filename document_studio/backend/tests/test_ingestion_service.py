"""Integration tests for the ingestion service with real format adapters.

Uses small deterministic fixtures generated at test time to prove:
- Digital PDF text extraction with page boundaries
- DOCX paragraph and table normalization
- XLSX sheet names, cell coordinates, and values
- Useful source locations for all formats
- Format detection and supported-format validation
- Safe filename handling and configurable size limits
- Unsupported-format and media-type mismatch errors
- Structured ``requires_ocr`` result for scanned/image-only PDFs
- JSON-compatible normalized serialization

No network calls, model keys, running server, OCR software, or
production paths are required.  All fixtures are generated in memory
using the parser libraries.
"""

from __future__ import annotations

import io
import json
import unittest

from document_studio.application.ingestion import (
    EmptyInputError,
    FileSizeLimitError,
    MediaTypeMismatchError,
    UnsafeFilenameError,
    UnsupportedFormatError,
)
from document_studio.domain.normalized import IngestionOutcome
from document_studio.infrastructure.ingestion_service import IngestionService


# ---------------------------------------------------------------------------
# Fixture generators (small, deterministic, in-memory)
# ---------------------------------------------------------------------------


def _make_pdf_with_text(text: str = "Hello, world!") -> bytes:
    """Create a minimal single-page PDF with embedded text using PyMuPDF."""
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=612, height=792)
    page.insert_text((72, 72), text, fontsize=12)
    data = doc.tobytes()
    doc.close()
    return data


def _make_pdf_image_only() -> bytes:
    """Create a minimal single-page PDF containing only an image (no text).

    This simulates a scanned document that requires OCR.
    """
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    # Insert a tiny 2x2 pixel PNG image (no text at all).
    # Minimal valid PNG: 2x2 red pixels.
    import struct
    import zlib

    def _minimal_png() -> bytes:
        """Generate a minimal 2x2 red PNG image."""
        width, height = 2, 2
        # Raw image data: filter byte 0 + RGB pixels
        raw_data = b""
        for _ in range(height):
            raw_data += b"\x00"  # filter byte
            for _ in range(width):
                raw_data += b"\xff\x00\x00"  # red pixel

        compressed = zlib.compress(raw_data)

        def _chunk(chunk_type: bytes, data: bytes) -> bytes:
            c = chunk_type + data
            crc = zlib.crc32(c) & 0xFFFFFFFF
            return struct.pack(">I", len(data)) + c + struct.pack(">I", crc)

        ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
        return (
            b"\x89PNG\r\n\x1a\n"
            + _chunk(b"IHDR", ihdr_data)
            + _chunk(b"IDAT", compressed)
            + _chunk(b"IEND", b"")
        )

    png_bytes = _minimal_png()
    page.insert_image(fitz.Rect(100, 100, 300, 300), stream=png_bytes)

    data = doc.tobytes()
    doc.close()
    return data


def _make_docx_with_content() -> bytes:
    """Create a minimal DOCX with paragraphs and a table."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Document Title")
    doc.add_paragraph("This is a test paragraph.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Total"
    table.cell(1, 1).text = "100"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx_with_sheets() -> bytes:
    """Create a minimal XLSX with two named sheets and cell values."""
    import openpyxl

    wb = openpyxl.Workbook()

    # First sheet (default).
    ws1 = wb.active
    ws1.title = "Revenue"
    ws1["A1"] = "Quarter"
    ws1["B1"] = "Amount"
    ws1["A2"] = "Q1"
    ws1["B2"] = 50000

    # Second sheet.
    ws2 = wb.create_sheet("Expenses")
    ws2["A1"] = "Category"
    ws2["B1"] = "Cost"
    ws2["A2"] = "Travel"
    ws2["B2"] = 1200

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestIngestionServiceValidation(unittest.TestCase):
    """Input validation before parsing."""

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_empty_input_raises(self) -> None:
        with self.assertRaises(EmptyInputError):
            self.svc.ingest(b"", "doc.pdf", "application/pdf")

    def test_unsafe_filename_slash_raises(self) -> None:
        with self.assertRaises(UnsafeFilenameError):
            self.svc.ingest(b"data", "path/file.pdf", "application/pdf")

    def test_unsafe_filename_backslash_raises(self) -> None:
        with self.assertRaises(UnsafeFilenameError):
            self.svc.ingest(b"data", "path\\file.pdf", "application/pdf")

    def test_unsafe_filename_empty_raises(self) -> None:
        with self.assertRaises(UnsafeFilenameError):
            self.svc.ingest(b"data", "", "application/pdf")

    def test_unsupported_extension_raises(self) -> None:
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(b"data", "image.png", "image/png")

    def test_unsupported_media_type_raises(self) -> None:
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(b"data", "file.pdf", "text/plain")

    def test_media_type_mismatch_raises(self) -> None:
        with self.assertRaises(MediaTypeMismatchError):
            docx_type = (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
            self.svc.ingest(b"data", "file.pdf", docx_type)

    def test_size_limit_default(self) -> None:
        """Default size limit does not reject small files."""
        pdf = _make_pdf_with_text("small")
        result = self.svc.ingest(pdf, "small.pdf", "application/pdf")
        self.assertEqual(result.outcome, IngestionOutcome.success)

    def test_size_limit_custom_exceeded(self) -> None:
        """Configurable size limit rejects oversized content."""
        pdf = _make_pdf_with_text("test")
        with self.assertRaises(FileSizeLimitError):
            self.svc.ingest(
                pdf, "big.pdf", "application/pdf", max_byte_size=10
            )

    def test_size_limit_exact_boundary(self) -> None:
        """Content exactly at the limit passes; one byte over fails."""
        pdf = _make_pdf_with_text("x")
        size = len(pdf)
        # Exactly at limit — should pass.
        result = self.svc.ingest(
            pdf, "exact.pdf", "application/pdf", max_byte_size=size
        )
        self.assertEqual(result.outcome, IngestionOutcome.success)
        # One byte over — should fail.
        with self.assertRaises(FileSizeLimitError):
            self.svc.ingest(
                pdf, "over.pdf", "application/pdf", max_byte_size=size - 1
            )


# ---------------------------------------------------------------------------
# PDF ingestion
# ---------------------------------------------------------------------------


class TestPdfIngestion(unittest.TestCase):
    """Digital PDF text extraction and scanned-PDF detection."""

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_digital_pdf_extracts_text(self) -> None:
        pdf = _make_pdf_with_text("Invoice #12345")
        result = self.svc.ingest(pdf, "invoice.pdf", "application/pdf")

        self.assertEqual(result.outcome, IngestionOutcome.success)
        self.assertIsNotNone(result.document)
        doc = result.document

        self.assertEqual(doc.media_type, "application/pdf")
        self.assertEqual(doc.original_filename, "invoice.pdf")
        self.assertGreater(len(doc.pages), 0)

        # At least one page must have text blocks.
        all_texts = []
        for page in doc.pages:
            for block in page.blocks:
                d = block.to_dict()
                if d["type"] == "text":
                    all_texts.append(d["text"])

        combined = " ".join(all_texts)
        self.assertIn("Invoice", combined)
        self.assertIn("12345", combined)

    def test_digital_pdf_preserves_page_boundaries(self) -> None:
        """Multi-page PDF produces multiple NormalizedPage objects."""
        import fitz

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i+1} content")
        data = doc.tobytes()
        doc.close()

        result = self.svc.ingest(data, "multi.pdf", "application/pdf")
        self.assertEqual(result.outcome, IngestionOutcome.success)
        self.assertEqual(len(result.document.pages), 3)

        for i, p in enumerate(result.document.pages):
            self.assertEqual(p.page_number, i + 1)

    def test_digital_pdf_source_locations(self) -> None:
        """Text blocks have useful page-based source locations."""
        pdf = _make_pdf_with_text("Located text")
        result = self.svc.ingest(pdf, "loc.pdf", "application/pdf")

        page = result.document.pages[0]
        block = page.blocks[0]
        loc = block.to_dict()["location"]

        self.assertEqual(loc["kind"], "page")
        self.assertEqual(loc["page_number"], 1)
        self.assertIsNotNone(loc["char_offset"])
        self.assertIsNotNone(loc["char_length"])
        self.assertGreater(loc["char_length"], 0)

    def test_scanned_pdf_requires_ocr(self) -> None:
        """Image-only PDF returns requires_ocr, not fake text."""
        pdf = _make_pdf_image_only()
        result = self.svc.ingest(pdf, "scanned.pdf", "application/pdf")

        self.assertEqual(result.outcome, IngestionOutcome.requires_ocr)
        self.assertIsNone(result.document)
        self.assertIsInstance(result.requires_ocr_reason, str)
        self.assertGreater(len(result.requires_ocr_reason), 0)

    def test_pdf_json_serialization(self) -> None:
        """Successful PDF result survives JSON round-trip."""
        pdf = _make_pdf_with_text("JSON test")
        result = self.svc.ingest(pdf, "json.pdf", "application/pdf")

        d = result.to_dict()
        raw = json.dumps(d)
        restored = json.loads(raw)
        self.assertEqual(d, restored)


# ---------------------------------------------------------------------------
# DOCX ingestion
# ---------------------------------------------------------------------------


class TestDocxIngestion(unittest.TestCase):
    """DOCX paragraph and table normalization."""

    DOCX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_docx_extracts_paragraphs(self) -> None:
        docx_data = _make_docx_with_content()
        result = self.svc.ingest(docx_data, "test.docx", self.DOCX_TYPE)

        self.assertEqual(result.outcome, IngestionOutcome.success)
        doc = result.document
        self.assertEqual(doc.media_type, self.DOCX_TYPE)
        self.assertGreater(len(doc.pages), 0)

        # Collect all text blocks.
        texts = []
        for page in doc.pages:
            for block in page.blocks:
                d = block.to_dict()
                if d["type"] == "text":
                    texts.append(d["text"])

        combined = " ".join(texts)
        self.assertIn("Document Title", combined)
        self.assertIn("test paragraph", combined)

    def test_docx_extracts_table(self) -> None:
        docx_data = _make_docx_with_content()
        result = self.svc.ingest(docx_data, "tbl.docx", self.DOCX_TYPE)

        # Find the table block.
        tables = []
        for page in result.document.pages:
            for block in page.blocks:
                d = block.to_dict()
                if d["type"] == "table":
                    tables.append(d)

        self.assertGreater(len(tables), 0)
        tbl = tables[0]
        self.assertEqual(tbl["row_count"], 2)
        self.assertEqual(tbl["col_count"], 2)

        cell_texts = {c["text"] for c in tbl["cells"]}
        self.assertIn("Name", cell_texts)
        self.assertIn("Value", cell_texts)
        self.assertIn("Total", cell_texts)
        self.assertIn("100", cell_texts)

    def test_docx_source_locations(self) -> None:
        docx_data = _make_docx_with_content()
        result = self.svc.ingest(docx_data, "loc.docx", self.DOCX_TYPE)

        page = result.document.pages[0]
        # First block should be a text block with page location.
        block = page.blocks[0]
        loc = block.to_dict()["location"]
        self.assertEqual(loc["kind"], "page")
        self.assertEqual(loc["page_number"], 1)

    def test_docx_json_serialization(self) -> None:
        docx_data = _make_docx_with_content()
        result = self.svc.ingest(docx_data, "json.docx", self.DOCX_TYPE)
        d = result.to_dict()
        raw = json.dumps(d)
        restored = json.loads(raw)
        self.assertEqual(d, restored)


# ---------------------------------------------------------------------------
# XLSX ingestion
# ---------------------------------------------------------------------------


class TestXlsxIngestion(unittest.TestCase):
    """XLSX sheet names, cell coordinates, and values."""

    XLSX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".spreadsheetml.sheet"
    )

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_xlsx_preserves_sheet_names(self) -> None:
        xlsx_data = _make_xlsx_with_sheets()
        result = self.svc.ingest(xlsx_data, "data.xlsx", self.XLSX_TYPE)

        self.assertEqual(result.outcome, IngestionOutcome.success)
        doc = result.document
        self.assertEqual(doc.media_type, self.XLSX_TYPE)
        self.assertGreater(len(doc.sheets), 0)

        sheet_names = [s.name for s in doc.sheets]
        self.assertIn("Revenue", sheet_names)
        self.assertIn("Expenses", sheet_names)

    def test_xlsx_preserves_cell_values(self) -> None:
        xlsx_data = _make_xlsx_with_sheets()
        result = self.svc.ingest(xlsx_data, "val.xlsx", self.XLSX_TYPE)

        # Find the Revenue sheet.
        revenue = None
        for sheet in result.document.sheets:
            if sheet.name == "Revenue":
                revenue = sheet
                break

        self.assertIsNotNone(revenue)
        self.assertGreater(len(revenue.blocks), 0)

        tbl = revenue.blocks[0].to_dict()
        self.assertEqual(tbl["type"], "table")

        cell_texts = {c["text"] for c in tbl["cells"]}
        self.assertIn("Quarter", cell_texts)
        self.assertIn("Amount", cell_texts)
        self.assertIn("Q1", cell_texts)
        self.assertIn("50000", cell_texts)

    def test_xlsx_preserves_cell_coordinates(self) -> None:
        xlsx_data = _make_xlsx_with_sheets()
        result = self.svc.ingest(xlsx_data, "coords.xlsx", self.XLSX_TYPE)

        sheet = result.document.sheets[0]
        tbl = sheet.blocks[0].to_dict()
        cells = tbl["cells"]

        # Verify row/col indices are present.
        for cell in cells:
            self.assertIn("row", cell)
            self.assertIn("col", cell)
            self.assertGreaterEqual(cell["row"], 0)
            self.assertGreaterEqual(cell["col"], 0)

    def test_xlsx_source_locations(self) -> None:
        xlsx_data = _make_xlsx_with_sheets()
        result = self.svc.ingest(xlsx_data, "loc.xlsx", self.XLSX_TYPE)

        sheet = result.document.sheets[0]
        tbl_block = sheet.blocks[0].to_dict()
        loc = tbl_block["location"]

        self.assertEqual(loc["kind"], "cell")
        self.assertIsNotNone(loc["sheet_name"])

    def test_xlsx_json_serialization(self) -> None:
        xlsx_data = _make_xlsx_with_sheets()
        result = self.svc.ingest(xlsx_data, "json.xlsx", self.XLSX_TYPE)
        d = result.to_dict()
        raw = json.dumps(d)
        restored = json.loads(raw)
        self.assertEqual(d, restored)


# ---------------------------------------------------------------------------
# Format detection via service
# ---------------------------------------------------------------------------


class TestFormatDetection(unittest.TestCase):
    """The service correctly routes formats based on extension + media type."""

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_pdf_detected(self) -> None:
        pdf = _make_pdf_with_text("detect")
        result = self.svc.ingest(pdf, "detect.pdf", "application/pdf")
        self.assertEqual(result.document.media_type, "application/pdf")

    def test_docx_detected(self) -> None:
        docx_data = _make_docx_with_content()
        mt = (
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        )
        result = self.svc.ingest(docx_data, "detect.docx", mt)
        self.assertEqual(result.document.media_type, mt)

    def test_xlsx_detected(self) -> None:
        xlsx_data = _make_xlsx_with_sheets()
        mt = (
            "application/vnd.openxmlformats-officedocument"
            ".spreadsheetml.sheet"
        )
        result = self.svc.ingest(xlsx_data, "detect.xlsx", mt)
        self.assertEqual(result.document.media_type, mt)

# ---------------------------------------------------------------------------
# Content-signature cross-labeling (mislabeled bytes)
# ---------------------------------------------------------------------------


class TestContentSignatureMismatch(unittest.TestCase):
    """Valid bytes deliberately mislabeled as a different supported format.

    These tests prove that the ingestion service rejects mislabeled
    content based on actual byte signatures and OOXML package contents,
    not just the filename extension and declared media type.
    """

    DOCX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )
    XLSX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".spreadsheetml.sheet"
    )

    def setUp(self) -> None:
        self.svc = IngestionService()
        # Generate real fixture bytes once per test.
        self.pdf_bytes = _make_pdf_with_text("real PDF")
        self.docx_bytes = _make_docx_with_content()
        self.xlsx_bytes = _make_xlsx_with_sheets()

    # -- OOXML cross-labeling (the exact defect scenario) ------------------

    def test_docx_bytes_labeled_as_xlsx_raises(self) -> None:
        """DOCX payload named report.xlsx with XLSX media type must fail."""
        with self.assertRaises(MediaTypeMismatchError) as ctx:
            self.svc.ingest(self.docx_bytes, "report.xlsx", self.XLSX_TYPE)
        self.assertIn("wordprocessingml", str(ctx.exception))

    def test_xlsx_bytes_labeled_as_docx_raises(self) -> None:
        """XLSX payload named data.docx with DOCX media type must fail."""
        with self.assertRaises(MediaTypeMismatchError) as ctx:
            self.svc.ingest(self.xlsx_bytes, "data.docx", self.DOCX_TYPE)
        self.assertIn("spreadsheetml", str(ctx.exception))

    # -- PDF vs OOXML cross-labeling --------------------------------------

    def test_pdf_bytes_labeled_as_docx_raises(self) -> None:
        """PDF bytes declared as DOCX must fail."""
        with self.assertRaises((MediaTypeMismatchError, UnsupportedFormatError)):
            self.svc.ingest(self.pdf_bytes, "fake.docx", self.DOCX_TYPE)

    def test_pdf_bytes_labeled_as_xlsx_raises(self) -> None:
        """PDF bytes declared as XLSX must fail."""
        with self.assertRaises((MediaTypeMismatchError, UnsupportedFormatError)):
            self.svc.ingest(self.pdf_bytes, "fake.xlsx", self.XLSX_TYPE)

    def test_docx_bytes_labeled_as_pdf_raises(self) -> None:
        """DOCX bytes declared as PDF must fail."""
        with self.assertRaises((MediaTypeMismatchError, UnsupportedFormatError)):
            self.svc.ingest(self.docx_bytes, "fake.pdf", "application/pdf")

    def test_xlsx_bytes_labeled_as_pdf_raises(self) -> None:
        """XLSX bytes declared as PDF must fail."""
        with self.assertRaises((MediaTypeMismatchError, UnsupportedFormatError)):
            self.svc.ingest(self.xlsx_bytes, "fake.pdf", "application/pdf")

    # -- Arbitrary / garbage bytes ----------------------------------------

    def test_garbage_bytes_as_pdf_raises(self) -> None:
        """Random bytes declared as PDF must fail."""
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(b"not a real file", "garbage.pdf", "application/pdf")

    def test_garbage_bytes_as_docx_raises(self) -> None:
        """Random bytes declared as DOCX must fail."""
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(
                b"not a real file", "garbage.docx", self.DOCX_TYPE
            )

    def test_garbage_bytes_as_xlsx_raises(self) -> None:
        """Random bytes declared as XLSX must fail."""
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(
                b"not a real file", "garbage.xlsx", self.XLSX_TYPE
            )

    # -- Correctly labeled bytes still work --------------------------------

    def test_correctly_labeled_pdf_passes(self) -> None:
        """Sanity: real PDF bytes with correct label still succeed."""
        result = self.svc.ingest(
            self.pdf_bytes, "real.pdf", "application/pdf"
        )
        self.assertEqual(result.outcome, IngestionOutcome.success)

    def test_correctly_labeled_docx_passes(self) -> None:
        """Sanity: real DOCX bytes with correct label still succeed."""
        result = self.svc.ingest(
            self.docx_bytes, "real.docx", self.DOCX_TYPE
        )
        self.assertEqual(result.outcome, IngestionOutcome.success)

    def test_correctly_labeled_xlsx_passes(self) -> None:
        """Sanity: real XLSX bytes with correct label still succeed."""
        result = self.svc.ingest(
            self.xlsx_bytes, "real.xlsx", self.XLSX_TYPE
        )
        self.assertEqual(result.outcome, IngestionOutcome.success)

# ---------------------------------------------------------------------------
# Regression: size limit enforced before ZIP inspection
# ---------------------------------------------------------------------------


class TestSizeLimitBeforeZipInspection(unittest.TestCase):
    """Oversized OOXML content must be rejected by FileSizeLimitError
    before any ZIP decompression or content-type inspection occurs.

    Regression for: ingestion_service.py previously ran
    verify_content_signature (which opens the ZIP) before enforcing
    the byte-size limit.
    """

    DOCX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )
    XLSX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".spreadsheetml.sheet"
    )

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_oversized_docx_rejected_before_inspection(self) -> None:
        """DOCX bytes exceeding max_byte_size raise FileSizeLimitError,
        not any ZIP-related or content-type error."""
        docx_data = _make_docx_with_content()
        # Set limit smaller than the actual DOCX.
        with self.assertRaises(FileSizeLimitError):
            self.svc.ingest(
                docx_data, "big.docx", self.DOCX_TYPE,
                max_byte_size=10,
            )

    def test_oversized_xlsx_rejected_before_inspection(self) -> None:
        """XLSX bytes exceeding max_byte_size raise FileSizeLimitError."""
        xlsx_data = _make_xlsx_with_sheets()
        with self.assertRaises(FileSizeLimitError):
            self.svc.ingest(
                xlsx_data, "big.xlsx", self.XLSX_TYPE,
                max_byte_size=10,
            )

    def test_oversized_pdf_rejected_before_inspection(self) -> None:
        """PDF bytes exceeding max_byte_size raise FileSizeLimitError."""
        pdf_data = _make_pdf_with_text("oversized test")
        with self.assertRaises(FileSizeLimitError):
            self.svc.ingest(
                pdf_data, "big.pdf", "application/pdf",
                max_byte_size=10,
            )


# ---------------------------------------------------------------------------
# Regression: macro-enabled OOXML rejection
# ---------------------------------------------------------------------------


def _make_macro_docm_bytes() -> bytes:
    """Create a minimal ZIP that mimics a DOCM package.

    Contains a [Content_Types].xml with the macro-enabled Word
    main content type, which should be rejected as unsupported.
    """
    import io
    import zipfile

    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/'
        'content-types">'
        '<Override PartName="/word/document.xml" ContentType='
        '"application/vnd.ms-word.document.macroEnabled.main+xml"/>'
        '</Types>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types_xml)
        zf.writestr("word/document.xml", "<document/>")
    return buf.getvalue()


def _make_macro_xlsm_bytes() -> bytes:
    """Create a minimal ZIP that mimics an XLSM package.

    Contains a [Content_Types].xml with the macro-enabled Excel
    main content type, which should be rejected as unsupported.
    """
    import io
    import zipfile

    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/'
        'content-types">'
        '<Override PartName="/xl/workbook.xml" ContentType='
        '"application/vnd.ms-excel.sheet.macroEnabled.main+xml"/>'
        '</Types>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types_xml)
        zf.writestr("xl/workbook.xml", "<workbook/>")
    return buf.getvalue()


class TestMacroEnabledOoxmlRejection(unittest.TestCase):
    """Macro-enabled OOXML packages (DOCM/XLSM) renamed to .docx/.xlsx
    must be rejected as UnsupportedFormatError.

    Regression for: ingestion.py previously used broad substring markers
    (wordprocessingml / spreadsheetml) that matched both standard and
    macro-enabled content types in [Content_Types].xml.
    """

    DOCX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )
    XLSX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".spreadsheetml.sheet"
    )

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_docm_renamed_to_docx_rejected(self) -> None:
        """A DOCM package renamed to .docx must not reach the parser."""
        docm_bytes = _make_macro_docm_bytes()
        with self.assertRaises(UnsupportedFormatError) as ctx:
            self.svc.ingest(docm_bytes, "report.docx", self.DOCX_TYPE)
        self.assertIn("Macro-enabled", str(ctx.exception))

    def test_xlsm_renamed_to_xlsx_rejected(self) -> None:
        """An XLSM package renamed to .xlsx must not reach the parser."""
        xlsm_bytes = _make_macro_xlsm_bytes()
        with self.assertRaises(UnsupportedFormatError) as ctx:
            self.svc.ingest(xlsm_bytes, "data.xlsx", self.XLSX_TYPE)
        self.assertIn("Macro-enabled", str(ctx.exception))

    def test_docm_renamed_to_pdf_rejected(self) -> None:
        """A DOCM package renamed to .pdf must be rejected."""
        docm_bytes = _make_macro_docm_bytes()
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(docm_bytes, "report.pdf", "application/pdf")

    def test_xlsm_renamed_to_pdf_rejected(self) -> None:
        """An XLSM package renamed to .pdf must be rejected."""
        xlsm_bytes = _make_macro_xlsm_bytes()
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(xlsm_bytes, "data.pdf", "application/pdf")

# ---------------------------------------------------------------------------
# Regression: malformed OOXML packages (missing main-part files)
# ---------------------------------------------------------------------------


def _make_malformed_docx_missing_main_part() -> bytes:
    """Create a ZIP with DOCX content-type declaration but no
    word/document.xml entry.

    This mimics a malformed or hand-crafted package that declares
    the correct content type but lacks the actual main-part file.
    """
    import io
    import zipfile

    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/'
        'content-types">'
        '<Override PartName="/word/document.xml" ContentType='
        '"application/vnd.openxmlformats-officedocument'
        '.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types_xml)
        # Deliberately omit word/document.xml
        zf.writestr("word/styles.xml", "<styles/>")
    return buf.getvalue()


def _make_malformed_xlsx_missing_main_part() -> bytes:
    """Create a ZIP with XLSX content-type declaration but no
    xl/workbook.xml entry.

    This mimics a malformed or hand-crafted package that declares
    the correct content type but lacks the actual main-part file.
    """
    import io
    import zipfile

    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/'
        'content-types">'
        '<Override PartName="/xl/workbook.xml" ContentType='
        '"application/vnd.openxmlformats-officedocument'
        '.spreadsheetml.sheet.main+xml"/>'
        '</Types>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types_xml)
        # Deliberately omit xl/workbook.xml
        zf.writestr("xl/styles.xml", "<styles/>")
    return buf.getvalue()


def _make_ooxml_with_type_text_only() -> bytes:
    """Create a ZIP where a DOCX type appears only inside an XML comment."""
    import io
    import zipfile

    content_types_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/'
        'content-types">'
        '<!-- application/vnd.openxmlformats-officedocument'
        '.wordprocessingml.document.main+xml -->'
        '</Types>'
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", content_types_xml)
        zf.writestr("word/document.xml", "<document/>")
    return buf.getvalue()


class TestMalformedOoxmlPackages(unittest.TestCase):
    """Malformed OOXML packages that declare the correct content type
    in [Content_Types].xml but are missing the actual main-part file
    must be rejected as UnsupportedFormatError.

    Regression for: _detect_ooxml_media_type previously accepted any
    ZIP with matching content-type text, allowing malformed packages
    to reach the parser adapter.
    """

    DOCX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )
    XLSX_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".spreadsheetml.sheet"
    )

    def setUp(self) -> None:
        self.svc = IngestionService()

    def test_docx_missing_main_part_rejected(self) -> None:
        """ZIP with DOCX content type but no word/document.xml must fail."""
        data = _make_malformed_docx_missing_main_part()
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(data, "bad.docx", self.DOCX_TYPE)

    def test_xlsx_missing_main_part_rejected(self) -> None:
        """ZIP with XLSX content type but no xl/workbook.xml must fail."""
        data = _make_malformed_xlsx_missing_main_part()
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(data, "bad.xlsx", self.XLSX_TYPE)

    def test_real_docx_still_accepted(self) -> None:
        """Sanity: a real DOCX (which has word/document.xml) still works."""
        docx_data = _make_docx_with_content()
        result = self.svc.ingest(docx_data, "good.docx", self.DOCX_TYPE)
        self.assertEqual(result.outcome, IngestionOutcome.success)

    def test_real_xlsx_still_accepted(self) -> None:
        """Sanity: a real XLSX (which has xl/workbook.xml) still works."""
        xlsx_data = _make_xlsx_with_sheets()
        result = self.svc.ingest(xlsx_data, "good.xlsx", self.XLSX_TYPE)
        self.assertEqual(result.outcome, IngestionOutcome.success)

    def test_content_type_text_outside_override_is_rejected(self) -> None:
        """Type-like text must be an exact ``Override`` declaration."""
        data = _make_ooxml_with_type_text_only()
        with self.assertRaises(UnsupportedFormatError):
            self.svc.ingest(data, "bad.docx", self.DOCX_TYPE)


if __name__ == "__main__":
    unittest.main()
